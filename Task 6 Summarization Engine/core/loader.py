from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document

def load_document(file_path: str):
    """
    Load PDF, TXT, or DOCX documents and return as LangChain Document list.
    """
    if file_path.lower().endswith(".pdf"):
        loader = PyPDFLoader(file_path)
        return loader.load()

    elif file_path.lower().endswith(".txt"):
        loader = TextLoader(file_path)
        return loader.load()

    elif file_path.lower().endswith(".docx"):
        # manually extract text from .docx
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        from langchain.schema import Document as LC_Document
        return [LC_Document(page_content=text, metadata={"source": file_path})]

    else:
        raise ValueError("Unsupported file format. Use PDF, TXT, or DOCX.")
